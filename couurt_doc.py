# encoding: utf-8
import requests
from pyquery import PyQuery as pq
import json
from pathlib import Path
import time
import random
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# session = requests.Session()
url = "http://www.iresearch.cn/"
trans_url = "http://api.t.sina.com.cn/short_url/shorten.json"
# page_range = [1,870]
relative_url = 'http://www.court.gov.cn'

header = {'User-Agent':'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 '
                       '(KHTML, like Gecko) Chrome/60.0.3112.78 Safari/537.36',
        'Accept':'text/html,application/xhtml+xml,application/xml;q=0.9,'
                 'image/webp,image/apng,*/*;q=0.8',
        'Accept-Encoding':'gzip, deflate',
        'Accept-Language':'zh-CN,zh;q=0.8,en-US;q=0.6,en;q=0.4',
        'Cache-Control':'no-cache',
        'Connection':'keep-alive'
          }

all_file_list = []


def get_all_page(i):
    url = 'http://www.court.gov.cn/paper/default/index/page/%d.html' % i
    content = requests.get(url,headers=header)
    doc = pq(content.text)
    all_content = doc('.list').find('.l').items()
    all = list(all_content)
    for c in all:
        title = c.children().find('.list_tit').text() # 文书名字
        href = c.children().find('.list_tit').children().attr('href') # 文书链接
        an_hao = c.children().find('.ah').text() # 文书案号
        public_date = c.children().find('.date').text() # 发布日期
        all_file_list.append({'title':title,'href':relative_url + href,'an_hao':an_hao,
                              'public_date':public_date})
        print({'title':title,'href':relative_url + href,'an_hao':an_hao,
                              'public_date':public_date})


def save_all():
    datas = json.dumps(all_file_list, ensure_ascii=False, indent=4)  # ensure_ascii：使用中文保存，缩进为4个空格
    with open('file.json', 'w+') as f:
        f.write(datas)


def doc():


    # 新建document对象
    document = Document()

    # 添加标题,level表示标题大小
    document.add_heading(u'标题', level=1)

    # 添加段落对象
    paragraph = document.add_paragraph("hello world".title())
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 右对齐

    # 在当前段落之前插入段落
    cc = paragraph.insert_paragraph_before(u"你好")
    cc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中


    # 添加换页
    # document.add_page_break()

    # 添加表格
    table = document.add_table(rows=2, cols=3)
    table.add_row()  # 添加行

    for row in table.rows:  # 遍历表格
        for cell in row.cells:
            cell.text = "fuck"
    cell = table.cell(0, 0)  # 单元格

    # # 添加图像并调整大小
    # document.add_picture("test.gif", width=Inches(1.0))

    # 样式
    paragraph = document.add_paragraph("Did i looking better?")
    paragraph.style = "ListBullet"

    # run
    paragraph = document.add_paragraph(u"你好你好年后")
    paragraph.paragraph_format.first_line_indent = Pt(20)  # 首行缩进2格
    run = paragraph.add_run("this is test run")  # 添加一句话，不会换行
    run.bold = True
    run.style = "Emphasis"  # 样式

    paragraph.add_run(u"你好你好年后年安徽.")
    paragraph.add_run(u"你好你好年后年安徽.")
    paragraph.add_run(u"你好你好年后年安徽.")

    # 保存文档
    document.save("test.docx")


if __name__ == '__main__':
    # for i in range(1,871):
    #     sleep = random.randint(7, 13)
    #     get_all_page(i)
    #     time.sleep(sleep)
    # save_all()
    doc()

